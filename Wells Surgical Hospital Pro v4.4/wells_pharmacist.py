import os
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    pass

class WellsPharmacist:
    def preparar_doses(self, lista_gabarito, dividir=True):
        """
        Versão 4.4: Gera Doses DOCX com limite cirúrgico de 3.200 linhas.
        Otimizado para a capacidade de processamento do Google Docs.
        """
        from wells_path_utils import obter_base_hospital
        base = obter_base_hospital()
        
        # Limpeza de resíduos de cirurgias anteriores
        for f in base.glob("GABARITO_DOSE_*.docx"): 
            try: os.remove(f)
            except: pass

        limite_linhas = 3200
        contador_dose = 1
        
        # Inicia a primeira dose
        doc = self.criar_novo_documento()
        linhas_na_dose_atual = 0

        for item in lista_gabarito:
            # Adiciona o bloco de tradução (Contexto + ID + Origem + TRADUÇÃO)
            p = doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(0)
            doc.add_paragraph("_" * 50).paragraph_format.space_after = Pt(6)
            
            linhas_na_dose_atual += 1

            # Se atingir o limite de 3200 linhas, sela a dose e abre uma nova
            if linhas_na_dose_atual >= limite_linhas:
                nome_final = base / f"GABARITO_DOSE_{contador_dose}.docx"
                doc.save(nome_final)
                self.registrar_log(f"Dose {contador_dose} selada com {linhas_na_dose_atual} linhas.")
                
                # Reseta para a próxima dose
                contador_dose += 1
                doc = self.criar_novo_documento()
                linhas_na_dose_atual = 0
                import gc; gc.collect() # Alívio de RAM para o Celeron

        # Salva a última dose (o restante das linhas)
        if linhas_na_dose_atual > 0:
            nome_final = base / f"GABARITO_DOSE_{contador_dose}.docx"
            doc.save(nome_final)
            self.registrar_log(f"Dose final {contador_dose} selada.")

        return base

    def criar_novo_documento(self):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Consolas'
        style.font.size = Pt(9) # Fonte menor para reduzir o peso do arquivo
        return doc

    def registrar_log(self, mensagem):
        from wells_path_utils import obter_base_hospital
        log_path = obter_base_hospital() / "hospital_log.txt"
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[Pharmacist v4.4] {mensagem}\n")
